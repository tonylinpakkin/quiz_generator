"""
File parsing utilities for extracting text from different file formats
"""
import io
import re
from typing import Tuple
import fitz  # PyMuPDF
from docx import Document

class FileParsingError(Exception):
    """Custom exception for file parsing errors"""
    pass

class FileParser:
    """Handles parsing of different file formats"""
    
    @staticmethod
    def extract_text_from_pdf(content: bytes) -> Tuple[str, int]:
        """Extract text from PDF content"""
        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=content, filetype="pdf")
            text_content = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                
                # Try multiple extraction methods to get readable content
                text = ""
                
                # Method 1: Standard text extraction
                standard_text = page.get_text()
                
                # Method 2: Dictionary-based extraction for better structure
                dict_text = ""
                try:
                    text_dict = page.get_text("dict")
                    for block in text_dict.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                line_text = ""
                                for span in line.get("spans", []):
                                    span_text = span.get("text", "").strip()
                                    if span_text:
                                        line_text += span_text + " "
                                if line_text.strip():
                                    dict_text += line_text.strip() + "\n"
                        elif block.get("type") == 0:  # Text block
                            # Handle direct text blocks
                            block_text = block.get("text", "").strip()
                            if block_text:
                                dict_text += block_text + "\n"
                except:
                    pass
                
                # Method 3: Try HTML extraction for complex layouts
                html_text = ""
                try:
                    html_content = page.get_text("html")
                    # Simple HTML tag removal for basic text extraction
                    html_text = re.sub(r'<[^>]+>', '', html_content)
                    html_text = re.sub(r'\s+', ' ', html_text).strip()
                except:
                    pass
                
                # Choose the method that gives the most readable content
                texts = [standard_text.strip(), dict_text.strip(), html_text.strip()]
                # Pick the longest non-metadata text
                best_text = ""
                for candidate in texts:
                    if (len(candidate) > len(best_text) and 
                        not candidate.startswith(('%PDF', '/Type', 'obj')) and
                        'FlateDecode' not in candidate):
                        best_text = candidate
                
                text = best_text if best_text else standard_text
                
                # Clean and validate the text
                if text.strip():
                    # Remove PDF artifacts and metadata
                    lines = text.split('\n')
                    clean_lines = []
                    for line in lines:
                        line = line.strip()
                        # Skip obvious PDF metadata/structure
                        if (line and 
                            not line.startswith(('%PDF', '/Type', '/Filter', 'obj', 'endobj', 'stream', 'endstream')) and
                            not line.replace(' ', '').replace('\t', '').startswith(('/', '%')) and
                            len(line) > 2):
                            clean_lines.append(line)
                    
                    if clean_lines:
                        text_content += "\n".join(clean_lines) + "\n\n"
            
            # Clean up technical artifacts while preserving content
            text_content = re.sub(r'/(?:Filter|FlateDecode|Length)\s+\d+', '', text_content)
            text_content = re.sub(r'%PDF-\d+\.\d+', '', text_content)
            text_content = re.sub(r'\[\d+\s+\d+\s+\d+\s+\d+\]', '', text_content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            # Validate extracted content
            if not text_content or text_content.isspace():
                raise FileParsingError("This PDF appears to contain images or scanned text that cannot be automatically extracted. Please try uploading a text-based PDF or convert your content to a Word document (.docx) or plain text file (.txt).")
            
            # Check if content appears to be just metadata
            metadata_lines = 0
            total_lines = 0
            for line in text_content.split('\n'):
                line = line.strip()
                if line:
                    total_lines += 1
                    if (line.startswith('/') or line.startswith('%') or 
                        'obj' in line or 'endobj' in line or 'stream' in line or 
                        'FlateDecode' in line or 'ICCBased' in line):
                        metadata_lines += 1
            
            # If most content is metadata, reject it
            if total_lines > 0 and metadata_lines / total_lines > 0.7:
                raise FileParsingError("This PDF contains mostly technical data rather than readable content. Your file might be image-based or use complex formatting. Please try uploading the content as a Word document (.docx) or plain text file (.txt) for better results.")
            
            pdf_document.close()
            
            # Clean up the text
            text_content = FileParser._clean_text(text_content)
            word_count = len(text_content.split())
            
            if not text_content.strip():
                raise FileParsingError("No readable text found in PDF")
            
            # Log the extracted content
            print("\n=== PDF Content Extraction Log ===")
            print(f"Word count: {word_count}")
            print("First 500 characters of extracted content:")
            print(text_content[:500] + "..." if len(text_content) > 500 else text_content)
            print("================================\n")
                
            return text_content, word_count
            
        except Exception as e:
            if isinstance(e, FileParsingError):
                raise
            raise FileParsingError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(content: bytes) -> Tuple[str, int]:
        """Extract text from DOCX content"""
        try:
            # Open document from bytes
            doc = Document(io.BytesIO(content))
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            # Clean up the text
            text_content = FileParser._clean_text(text_content)
            word_count = len(text_content.split())
            
            if not text_content.strip():
                raise FileParsingError("No readable text found in DOCX")
                
            return text_content, word_count
            
        except Exception as e:
            if isinstance(e, FileParsingError):
                raise
            raise FileParsingError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(content: bytes) -> Tuple[str, int]:
        """Extract text from TXT content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text_content = None
            
            for encoding in encodings:
                try:
                    text_content = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise FileParsingError("Unable to decode text file with supported encodings")
            
            # Clean up the text
            text_content = FileParser._clean_text(text_content)
            word_count = len(text_content.split())
            
            if not text_content.strip():
                raise FileParsingError("Text file appears to be empty")
                
            return text_content, word_count
            
        except Exception as e:
            if isinstance(e, FileParsingError):
                raise
            raise FileParsingError(f"Failed to parse TXT: {str(e)}")
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace and line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Normalize paragraph breaks
        text = re.sub(r'[ \t]+', ' ', text)      # Normalize spaces
        text = text.strip()
        
        # Remove control characters except newlines and tabs
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        return text
    
    @staticmethod
    def parse_file(filename: str, content: bytes) -> Tuple[str, int]:
        """Parse file based on extension and return text content and word count"""
        file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if file_extension == 'pdf':
            return FileParser.extract_text_from_pdf(content)
        elif file_extension == 'docx':
            return FileParser.extract_text_from_docx(content)
        elif file_extension == 'txt':
            return FileParser.extract_text_from_txt(content)
        else:
            raise FileParsingError(f"Unsupported file format: {file_extension}")

def validate_file_type(filename: str) -> bool:
    """Validate if file type is supported"""
    supported_extensions = {'pdf', 'docx', 'txt'}
    file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
    return file_extension in supported_extensions

def get_file_type(filename: str) -> str:
    """Get file type from filename"""
    file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
    return file_extension
